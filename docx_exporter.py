from docx import Document
import io

def gerar_docx_relatorio(texto, nome_arquivo="relatorio_pcdt.docx"):
    doc = Document()
    doc.add_heading("Relatório PCDT - Análise de Exames", level=1)

    for linha in texto.split("\n"):
        doc.add_paragraph(linha)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer