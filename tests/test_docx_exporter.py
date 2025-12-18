"""
Tests for docx_exporter.py

Tests DOCX report generation functionality.
"""
import pytest
import io
from docx_exporter import gerar_docx_relatorio
from docx import Document


class TestGerarDocxRelatorio:
    """Tests for DOCX report generation"""

    @pytest.mark.unit
    def test_gerar_docx_relatorio_returns_buffer(self):
        """Should return a BytesIO buffer"""
        texto = "Paciente: João Silva\nIdade: 65"
        result = gerar_docx_relatorio(texto)

        assert isinstance(result, io.BytesIO)
        assert result.tell() == 0  # Buffer should be at position 0

    @pytest.mark.unit
    def test_gerar_docx_relatorio_creates_valid_docx(self):
        """Should create a valid DOCX file"""
        texto = "Paciente: João Silva\nIdade: 65"
        result = gerar_docx_relatorio(texto)

        # Read the buffer and verify it's a valid DOCX
        docx_content = result.read()

        # DOCX files are ZIP archives, should start with PK signature
        assert docx_content.startswith(b'PK')
        assert len(docx_content) > 0

    @pytest.mark.unit
    def test_gerar_docx_relatorio_has_heading(self):
        """Should include the standard heading"""
        texto = "Test content"
        result = gerar_docx_relatorio(texto)

        # Parse the DOCX to verify content
        result.seek(0)
        doc = Document(result)

        # First paragraph should be the heading
        assert len(doc.paragraphs) > 0
        # The heading text should be present in the document
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Relatório PCDT" in full_text or "Análise de Exames" in full_text

    @pytest.mark.unit
    def test_gerar_docx_relatorio_includes_content(self):
        """Should include the provided text content"""
        texto = "Paciente: João Silva\nIdade: 65\nDiagnóstico: Normal"
        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        # Extract all text from document
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Content should be present
        assert "Paciente: João Silva" in full_text
        assert "Idade: 65" in full_text
        assert "Diagnóstico: Normal" in full_text

    @pytest.mark.unit
    def test_gerar_docx_relatorio_multiline_text(self):
        """Should preserve multiple lines as separate paragraphs"""
        texto = "Line 1\nLine 2\nLine 3"
        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Line 1" in full_text
        assert "Line 2" in full_text
        assert "Line 3" in full_text

    @pytest.mark.unit
    def test_gerar_docx_relatorio_empty_text(self):
        """Should handle empty text"""
        texto = ""
        result = gerar_docx_relatorio(texto)

        assert result is not None
        docx_content = result.read()
        assert docx_content.startswith(b'PK')

    @pytest.mark.unit
    def test_gerar_docx_relatorio_special_characters(self):
        """Should handle Portuguese special characters"""
        texto = "Cálcio: 9.2 mg/dL\nFósforo: 5.5 mg/dL\nDiálise"
        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Cálcio" in full_text
        assert "Fósforo" in full_text
        assert "Diálise" in full_text

    @pytest.mark.unit
    def test_gerar_docx_relatorio_default_filename(self):
        """Should work with default filename parameter"""
        texto = "Test"
        result = gerar_docx_relatorio(texto)

        assert result is not None
        docx_content = result.read()
        assert len(docx_content) > 0

    @pytest.mark.unit
    def test_gerar_docx_relatorio_custom_filename(self):
        """Should work with custom filename parameter"""
        texto = "Test"
        custom_name = "custom_report.docx"
        result = gerar_docx_relatorio(texto, nome_arquivo=custom_name)

        assert result is not None
        docx_content = result.read()
        assert docx_content.startswith(b'PK')

    @pytest.mark.unit
    def test_gerar_docx_relatorio_buffer_position(self):
        """Should reset buffer position to start"""
        texto = "Test content"
        result = gerar_docx_relatorio(texto)

        # Buffer should be at position 0
        assert result.tell() == 0

        # Should be readable
        content = result.read()
        assert len(content) > 0

    @pytest.mark.unit
    def test_gerar_docx_relatorio_complete_report(self):
        """Should generate DOCX for a complete clinical report"""
        texto = """Paciente: João Silva Santos
Idade: 65
Modalidade: Hemodiálise

Diagnósticos prováveis:
- Anemia da DRC
- Hiperparatireoidismo secundário

Condutas sugeridas:
- Iniciar alfaepoetina e avaliar ferro sérico.
- Reposição de ferro (ex: sacarato férrico).

Evolução clínica automática:
Paciente em diálise com alterações laboratoriais."""

        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Verify key sections are present
        assert "Paciente: João Silva Santos" in full_text
        assert "Anemia da DRC" in full_text
        assert "Hiperparatireoidismo secundário" in full_text
        assert "alfaepoetina" in full_text

    @pytest.mark.unit
    def test_gerar_docx_relatorio_long_text(self):
        """Should handle long text content"""
        lines = [f"Line {i}: Patient exam data and results" for i in range(100)]
        texto = "\n".join(lines)

        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        # Should have many paragraphs (heading + 100 lines)
        assert len(doc.paragraphs) > 50

    @pytest.mark.unit
    def test_gerar_docx_relatorio_returns_new_buffer_each_time(self):
        """Should return a new buffer for each call"""
        texto = "Test"

        result1 = gerar_docx_relatorio(texto)
        result2 = gerar_docx_relatorio(texto)

        # Should be different buffer objects
        assert result1 is not result2
        assert id(result1) != id(result2)

    @pytest.mark.unit
    def test_gerar_docx_relatorio_paragraph_count(self):
        """Should create correct number of paragraphs"""
        texto = "Line 1\nLine 2\nLine 3"
        result = gerar_docx_relatorio(texto)

        result.seek(0)
        doc = Document(result)

        # Should have: 1 heading + 3 content lines = 4+ paragraphs
        # (heading is added, plus one paragraph per line in the loop)
        assert len(doc.paragraphs) >= 4
